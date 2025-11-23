import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit
import textwrap

class ReportGenerator:
    def __init__(self, output_dir="reports"):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def generate_docx(self, content: str, filename: str) -> str:
        """Generates a DOCX report."""
        doc = Document()
        doc.add_heading('Research Report', 0)
        
        # Split content by newlines and add as paragraphs
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
                
        filepath = os.path.join(self.output_dir, f"{filename}.docx")
        doc.save(filepath)
        return filepath

    def generate_pdf(self, content: str, filename: str) -> str:
        """Generates a PDF report."""
        filepath = os.path.join(self.output_dir, f"{filename}.pdf")
        c = canvas.Canvas(filepath, pagesize=letter)
        width, height = letter
        
        # Title
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, height - 50, "Research Report")
        
        # Content
        c.setFont("Helvetica", 12)
        y_position = height - 80
        margin = 50
        max_width = width - 2 * margin
        
        for line in content.split('\n'):
            # Wrap text
            wrapped_lines = simpleSplit(line, "Helvetica", 12, max_width)
            for wrapped_line in wrapped_lines:
                if y_position < 50:
                    c.showPage()
                    c.setFont("Helvetica", 12)
                    y_position = height - 50
                
                c.drawString(margin, y_position, wrapped_line)
                y_position -= 15
            
            # Add extra space between paragraphs
            y_position -= 10
            
        c.save()
        return filepath
